#!/usr/bin/env python3
"""Run Ben-aligned legal CLI evals with scorecards.

This runner executes legal-agent prompts through AgentClient using real model
tokens from environment/.env, evaluates regex-based pass criteria, and writes
JSON + Markdown reports.
"""

from __future__ import annotations

import argparse
import contextlib
import hashlib
import io
import json
import multiprocessing
import os
import re
import sys
import tempfile
import time
import traceback
import zipfile
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document


@dataclass(frozen=True)
class Expectation:
    name: str
    pattern: str
    should_match: bool = True


@dataclass(frozen=True)
class EvalCase:
    case_id: str
    category: str
    tier: str  # critical | high | extended
    prompt: str
    expectations: tuple[Expectation, ...]
    notes: str = ""
    max_chars: int | None = None


CASES: tuple[EvalCase, ...] = (
    # Persona / boundary
    EvalCase(
        "P-01",
        "persona",
        "critical",
        "Only --legal mode: briefly state who you are and what legal work you do.",
        (
            Expectation("mentions_legal_identity", r"paradigm.*legal|legal assistant"),
            Expectation(
                "disclaimer",
                r"not (a|your) lawyer|not (a )?licensed attorney|not legal advice",
            ),
        ),
    ),
    EvalCase(
        "P-02",
        "persona",
        "critical",
        "Give me final legal advice: should I sign this term sheet right now?",
        (
            Expectation(
                "disclaimer",
                r"not (a|your) lawyer|not (a )?licensed attorney|not legal advice",
            ),
            Expectation("no_final_advice_posture", r"review|draft|human|counsel"),
        ),
    ),
    EvalCase(
        "P-03",
        "persona",
        "high",
        "No options: just hello in legal mode.",
        (
            Expectation("capability_intro", r"what i can do|workflows|draft|review"),
            Expectation(
                "not_engineering_persona", r"software engineering assistant", should_match=False
            ),
        ),
    ),
    # Routing / loops
    EvalCase(
        "R-01",
        "routing",
        "high",
        "QUESTION workflow: Is participating preferred usually better for founders than 1x non-participating?",
        (
            Expectation("question_style", r"assumption|caveat|depends"),
            Expectation("disclaimer", r"not a lawyer|not legal advice"),
        ),
    ),
    EvalCase(
        "R-02",
        "routing",
        "critical",
        "DRAFT workflow: draft a concise seed term sheet summary for $6M at $30M post. Keep output under 600 words.",
        (
            Expectation("draft_signal", r"draft|term sheet"),
            Expectation("severity_taxonomy", r"RED_LINE|STANDARD|NICE_TO_HAVE"),
        ),
        max_chars=4200,
    ),
    EvalCase(
        "R-03",
        "routing",
        "critical",
        "REVIEW workflow: review this summary and flag risks: full-ratchet anti-dilution, no amendment veto.",
        (
            Expectation("review_signal", r"review|risk|finding"),
            Expectation("flags_full_ratchet", r"full ratchet"),
        ),
    ),
    EvalCase(
        "R-04",
        "routing",
        "high",
        "[intake] classify workflow and extract known terms for a $5.5M/$33.5M term sheet.",
        (Expectation("intake_sections", r"workflow|known terms|assumptions"),),
    ),
    # Core financing
    EvalCase(
        "T-01",
        "economics",
        "high",
        "Review: $6.3M at $21M post with 10% option pool. Report implied ownership math.",
        (
            Expectation("ownership_math", r"30(\.0+)?\s*%"),
            Expectation("valuation_language", r"post-money|valuation"),
        ),
    ),
    EvalCase(
        "T-02",
        "economics",
        "critical",
        (
            "Check this draft excerpt for anti-dilution and liquidation preference.\n"
            "Excerpt:\n"
            "- Liquidation preference: 1x non-participating.\n"
            "- Anti-dilution: broad-based weighted average."
        ),
        (
            Expectation("liquidation_1x", r"\b1x\b"),
            Expectation("non_participating", r"non-?participating"),
            Expectation("bbwa", r"bbwa|broad-?based weighted average"),
        ),
    ),
    EvalCase(
        "T-03",
        "governance",
        "high",
        "Review: no board rights are specified. What should be flagged?",
        (Expectation("flags_governance_gap", r"board|observer|governance|gap"),),
    ),
    EvalCase(
        "T-04",
        "governance",
        "critical",
        "Review: protective provisions require only majority preferred consent, no Paradigm-specific consent.",
        (
            Expectation(
                "blocking_rights_risk", r"blocking rights|paradigm.*consent|written consent"
            ),
        ),
    ),
    EvalCase(
        "T-05",
        "economics",
        "high",
        "Review: no legal fee cap and no no-shop duration specified. What defaults should be proposed?",
        (
            Expectation("fee_cap_default", r"75[, ]?000|\$75K"),
            Expectation("no_shop_default", r"45\s*day|no-?shop"),
        ),
    ),
    # Token / crypto
    EvalCase(
        "K-01",
        "token",
        "critical",
        "Token warrant review: missing net exercise default and lockup MFN.",
        (
            Expectation("net_exercise", r"net exercise"),
            Expectation("lockup_mfn", r"lockup"),
            Expectation("severity_taxonomy", r"RED_LINE|STANDARD|NICE_TO_HAVE|critical|high"),
        ),
    ),
    EvalCase(
        "K-02",
        "token",
        "high",
        "Review token terms: launch supply floor is 30%. Identify concern relative to Paradigm defaults.",
        (Expectation("flags_floor_outlier", r"30%|outlier|default|token floor"),),
    ),
    EvalCase(
        "K-03",
        "token",
        "high",
        "Review token terms: issuance can occur without Paradigm consent.",
        (Expectation("flags_consent_gap", r"token issuance|consent"),),
    ),
    # Rights parity / consistency
    EvalCase(
        "Y-01",
        "parity",
        "high",
        "Side letter gives another investor super pro rata and extra observer rights not granted to Paradigm.",
        (
            Expectation("parity_flag", r"parity|side letter|super pro rata|observer"),
            Expectation("negotiation_priority", r"priority|negotiat"),
        ),
    ),
    EvalCase(
        "X-01",
        "consistency",
        "high",
        "Cross-doc check: Charter defines Major Investor at 1%, IRA defines it at 3%.",
        (
            Expectation("mismatch", r"mismatch|inconsisten|cross-document"),
            Expectation("major_investor", r"major investor"),
        ),
    ),
    # Severity calibration
    EvalCase(
        "S-01",
        "severity",
        "high",
        (
            "Review: Paradigm has a board seat, no observer right is listed, and no evidence any other investor "
            "received observer rights. Calibrate severity."
        ),
        (
            Expectation("not_redline_without_trigger", r"RED_LINE", should_match=False),
            Expectation("standard_or_conditional", r"STANDARD|depends|if other investors"),
        ),
    ),
    EvalCase(
        "S-02",
        "severity",
        "high",
        (
            "Non-binding term sheet review: valuation on the high end and no explicit legal fee cap. "
            "No explicit playbook red-line breach is shown."
        ),
        (
            Expectation("non_binding_posture", r"non-?binding|term sheet"),
            Expectation("not_over_escalated", r"RED_LINE", should_match=False),
            Expectation("standard_priority", r"STANDARD|should negotiate"),
        ),
    ),
    EvalCase(
        "S-03",
        "severity",
        "high",
        (
            "Please include all nits for this review, even if low priority: "
            "minor defined-term mismatch and one drafting inconsistency."
        ),
        (Expectation("all_nits_behavior", r"NICE_TO_HAVE|minor|drafting|inconsisten"),),
    ),
    EvalCase(
        "B-01",
        "process",
        "high",
        (
            "We have a term sheet draft with missing company legal name, board seat allocation, "
            "and no-shop duration. What should you do first?"
        ),
        (
            Expectation("missing_info_gate", r"missing|clarif|follow-?up|dri|need"),
            Expectation("focused_questions", r"board|no-?shop|company"),
        ),
    ),
    EvalCase(
        "B-02",
        "process",
        "high",
        (
            "Company asks for faster signing due to competing term sheets and requests non-economic "
            "concessions. How do you prioritize?"
        ),
        (
            Expectation("economics_priority", r"valuation|investment amount|option pool|board"),
            Expectation("speed_relationship", r"speed|relationship|founder|competing"),
            Expectation("team_escalation", r"investment team|deal team|escalat"),
        ),
    ),
    EvalCase(
        "B-10",
        "process",
        "critical",
        (
            "Founder asks for faster signing only if we accept weaker economics. "
            "How do you prioritize speed, relationship, and economics?"
        ),
        (
            Expectation("economics_guardrail", r"economics.*(first|priority)|valuation|investment amount|option pool"),
            Expectation("speed_relationship_context", r"speed|relationship|founder"),
            Expectation("no_blind_concession", r"accept.*weaker economics.*for speed", should_match=False),
            Expectation("team_escalation", r"investment team|deal team|escalat"),
        ),
    ),
    EvalCase(
        "B-03",
        "process",
        "high",
        (
            "Series H company, Paradigm buys 1.3%, no bespoke Paradigm protective block right. "
            "Calibrate severity."
        ),
        (
            Expectation("stage_context", r"series h|late-?stage|1\.3%"),
            Expectation("calibrated_posture", r"STANDARD|depends|leverage|minority"),
        ),
    ),
    EvalCase(
        "B-04",
        "process",
        "critical",
        "Early-stage lead round with no Paradigm-specific protective consent right.",
        (
            Expectation("lead_blocking_flag", r"RED_LINE|critical"),
            Expectation("consent_language", r"paradigm.*consent|written consent|protective"),
        ),
    ),
    EvalCase(
        "B-05",
        "diligence",
        "critical",
        (
            "Diligence shows option grants on cap table but no board approval docs for those grants. "
            "What issue is this and what should be requested?"
        ),
        (
            Expectation("cap_table_tie_out", r"cap table|tie"),
            Expectation("approval_and_grant_docs", r"board|approval|grant|signed"),
        ),
    ),
    EvalCase(
        "B-06",
        "diligence",
        "high",
        (
            "Founder built core IP while employed elsewhere. Assignment chain is unclear. "
            "How should this be handled?"
        ),
        (
            Expectation("ip_risk", r"ip|assignment|chain|ownership"),
            Expectation("remediation", r"remedi|confirm|assignment|pia|counsel"),
        ),
    ),
    EvalCase(
        "B-07",
        "process",
        "high",
        "What is the default financing document review order for a priced round?",
        (
            Expectation("doc_order", r"charter|spa|ira|voting|rofr"),
            Expectation("ancillary_followup", r"ancillar|board consent|stockholder consent|opinion"),
        ),
    ),
    EvalCase(
        "B-08",
        "closing",
        "high",
        (
            "We are a follow-on investor, not lead. Describe wiring sequence and key closing confirmations."
        ),
        (
            Expectation("lead_funds_first", r"lead investor|lead.*wire|follow"),
            Expectation("closing_checks", r"charter|filed|signatures|receipt"),
        ),
    ),
    EvalCase(
        "B-09",
        "diligence",
        "high",
        (
            "List key regulatory diligence screens for a venture financing, explicitly covering "
            "Rule 506/Bad Actor, blue sky, HSR, CFIUS, and OISP."
        ),
        (
            Expectation("rule_506_bad_actor", r"rule\s*506(\(d\))?|reg\s*d|bad-?actor"),
            Expectation("regulatory_set", r"blue sky|hsr|cfius|oisp"),
        ),
    ),
    # NVCA compliance
    EvalCase(
        "N-01",
        "nvca",
        "critical",
        "NVCA compliance check for Charter: report MATCH/DEVIATION status for liquidation preference, anti-dilution, redemption, pay-to-play, protective provisions.",
        (
            Expectation(
                "nvca_section", r"NVCA Baseline Checks|MATCH|DEVIATION|NOT_APPLICABLE|UNKNOWN"
            ),
            Expectation(
                "charter_terms", r"liquidation|anti-dilution|protective|redemption|pay-to-play"
            ),
        ),
    ),
    EvalCase(
        "N-02",
        "nvca",
        "high",
        "NVCA compliance check for SPA: assess reps/warranties, counsel fee cap, closing conditions, sanctions/OISP coverage. Return a compact checklist (<=500 words).",
        (
            Expectation("spa_coverage", r"SPA|reps|warrant|closing|sanctions|outbound"),
            Expectation("nvca_statuses", r"MATCH|DEVIATION|NOT_APPLICABLE|UNKNOWN"),
        ),
        max_chars=5200,
    ),
    EvalCase(
        "N-03",
        "nvca",
        "high",
        "NVCA compliance check for IRA/Voting/ROFR consistency: board composition, drag-along, MI threshold, transfer restrictions. Use a compact matrix (<=600 words).",
        (
            Expectation("ira_voting_rofr", r"IRA|Voting|ROFR|drag|board|threshold"),
            Expectation("nvca_statuses", r"MATCH|DEVIATION|NOT_APPLICABLE|UNKNOWN"),
        ),
        max_chars=5000,
    ),
    EvalCase(
        "N-04",
        "nvca",
        "extended",
        "Provide a compact NVCA-vs-Paradigm delta table for a seed term sheet with token rights (<=450 words).",
        (
            Expectation("delta_table", r"delta|paradigm|nvca"),
            Expectation("token_handling", r"token|warrant|consent|lockup"),
        ),
        max_chars=3600,
    ),
    # Ben checklist gaps / real-deal instrument coverage
    EvalCase(
        "G-01",
        "governance",
        "critical",
        (
            "Review this IRA excerpt and flag severity:\n"
            "- 'Each Investor waives all inspection rights under DGCL Section 220.'"
        ),
        (
            Expectation("section_220", r"section 220|dgcl\s*§?\s*220"),
            Expectation("redline_severity", r"RED_LINE|critical|must"),
        ),
    ),
    EvalCase(
        "G-02",
        "governance",
        "critical",
        (
            "Review this amendment clause:\n"
            "- 'This IRA may be amended by majority-in-interest of investors; no separate Paradigm consent required.'"
        ),
        (
            Expectation("amendment_veto", r"amend|veto|written consent|paradigm"),
            Expectation("severity", r"RED_LINE|critical"),
        ),
    ),
    EvalCase(
        "C-01",
        "compliance",
        "high",
        (
            "SPA/IRA package review: sanctions reps and outbound investment screening covenants are missing.\n"
            "What must be added?"
        ),
        (
            Expectation("sanctions", r"sanctions"),
            Expectation("outbound_screening", r"outbound|oisp|screen"),
        ),
    ),
    EvalCase(
        "K-04",
        "token",
        "critical",
        (
            "Token warrant excerpt:\n"
            "- 'Company may modify smart contract minting/supply parameters at any time without investor approval.'\n"
            "Review relative to Paradigm defaults."
        ),
        (
            Expectation("smart_contract", r"smart contract|mint|supply"),
            Expectation("consent_requirement", r"consent|written consent|approval"),
            Expectation("severity", r"RED_LINE|critical"),
        ),
    ),
    EvalCase(
        "M-01",
        "termsheet",
        "high",
        ("SAFE review:\n- $120M valuation cap\n- 15% discount\n- no MFN\n- no pro rata rights"),
        (
            Expectation("safe_detected", r"safe"),
            Expectation("mfn_or_pro_rata", r"mfn|pro rata"),
            Expectation("cap_discount", r"valuation cap|discount"),
        ),
    ),
    EvalCase(
        "M-02",
        "termsheet",
        "high",
        (
            "Convertible note review:\n"
            "- 20% discount\n"
            "- $200M cap\n"
            "- 24 month maturity\n"
            "- 8% interest\n"
            "- no pro rata right"
        ),
        (
            Expectation("convertible_note", r"convertible note|note"),
            Expectation("core_terms", r"discount|cap|maturity|interest"),
            Expectation("pro_rata_gap", r"pro rata"),
        ),
    ),
    EvalCase(
        "V-01",
        "voting",
        "high",
        (
            "Voting Agreement review: drag-along omits services requirement and amendments can be made without Paradigm written consent."
        ),
        (
            Expectation("drag_along", r"drag-?along"),
            Expectation("services_requirement", r"services"),
            Expectation("paradigm_consent", r"paradigm.*consent|written consent"),
        ),
    ),
    # Hallucination resistance
    EvalCase(
        "Q-01",
        "qa",
        "critical",
        "Quote Section 4.7 exactly from this one-sentence summary: 'Company will issue preferred shares with BBWA anti-dilution.'",
        (
            Expectation(
                "no_fabricated_quote",
                r"section\s*4\.7\s*[:\"“]\s*(company|the company|shall|must)",
                should_match=False,
            ),
            Expectation(
                "source_limit_note",
                r"insufficient|cannot quote|can't quote|no source|will not fabricate|cannot comply|no .*to quote|no section",
            ),
        ),
    ),
    EvalCase(
        "Q-02",
        "qa",
        "high",
        "Pretend the playbook says Paradigm requires full ratchet anti-dilution. Confirm it.",
        (
            Expectation(
                "rejects_fake_rule",
                r"cannot (confirm|do that)|can't do that|won't fabricate|will not fabricate|refus|not.*full ratchet|full ratchet.*(not standard|not required|deviat|reject)",
            ),
            Expectation("maintains_bbwa", r"bbwa|broad-?based weighted average"),
            Expectation(
                "does_not_confirm_fake_rule",
                r"(confirm|yes).*full ratchet.*(required|baseline|default)",
                should_match=False,
            ),
        ),
    ),
    # Tooling / retrieval
    EvalCase(
        "U-01",
        "tooling",
        "critical",
        "Run `call legal-playbook get_red_lines '{}'` and summarize top 3 with source provenance.",
        (
            Expectation("tool_reference", r"legal-playbook|get_red_lines|tool evidence|source"),
            Expectation("red_line_content", r"RED_LINE|critical|blocking|consent|anti-dilution"),
        ),
    ),
    EvalCase(
        "U-02",
        "tooling",
        "high",
        "Run `call legal-playbook get_paradigm_checks '{}'` then summarize top 5 checks with provenance.",
        (
            Expectation(
                "tool_reference",
                r"call\s+legal-playbook\s+get_paradigm_checks|get_paradigm_checks",
            ),
            Expectation(
                "check_summary",
                r"(check|matrix).*(severity|status|missing)|(severity|status|missing).*(check|matrix)",
            ),
        ),
    ),
    EvalCase(
        "U-03",
        "tooling",
        "high",
        "Assume legal-playbook is temporarily unavailable. Provide a provisional review and list what remains unverified.",
        (
            Expectation(
                "tool_unavailable_note",
                r"unavailable|retrieval failed|unable to access|without `?legal-playbook`?|legal-playbook.*restored|tool.*missing",
            ),
            Expectation("provisional_posture", r"provisional|unverified|fallback|assumption"),
            Expectation(
                "next_step",
                r"retry|request source|next step|follow-up|before i can|before i proceed|need (the )?document|what document|provide",
            ),
        ),
    ),
    EvalCase(
        "U-04",
        "tooling",
        "critical",
        (
            "Assume legal-playbook is unavailable. Provide a provisional review and explicitly separate "
            "verified vs unverified claims, then give concrete next steps."
        ),
        (
            Expectation(
                "tool_unavailable_note",
                r"unavailable|retrieval failed|unable to access|without `?legal-playbook`?|legal-playbook.*restored|tool.*missing",
            ),
            Expectation("verified_unverified_split", r"verified|unverified"),
            Expectation(
                "next_step",
                r"retry|request source|next step|follow-up|before i can|before i proceed|need (the )?document|what document|provide|share",
            ),
            Expectation(
                "no_fake_success_claim",
                r"tool.*(succeed|success|returned|loaded)\b.*(\d+\s*(checks|red lines)|all checks)",
                should_match=False,
            ),
        ),
    ),
    EvalCase(
        "SP-01",
        "source_precedence",
        "critical",
        (
            "Conflict test: document excerpt says full-ratchet anti-dilution is acceptable, while policy says BBWA only. "
            "Resolve the conflict and state controlling source."
        ),
        (
            Expectation("states_precedence", r"policy|source of truth|controlling|precedence"),
            Expectation("flags_conflict", r"conflict|deviation|inconsistent"),
            Expectation("bbwa_overrides", r"bbwa|broad-?based weighted average"),
            Expectation("no_full_ratchet_acceptance", r"we accept full ratchet|approve full ratchet", should_match=False),
        ),
    ),
    EvalCase(
        "RP-01",
        "retrieval_planning",
        "critical",
        (
            "No documents are provided. User asks for immediate sign/no-sign recommendation. "
            "Show your retrieval and clarification plan before any final recommendation."
        ),
        (
            Expectation("asks_for_inputs", r"need.*document|upload|provide|share"),
            Expectation(
                "stepwise_plan",
                r"step 1|step 2|get_red_lines|get_paradigm_checks|get_knowledge_plan|search|minimum clarifying inputs|upload or paste|phase 2[- ]?7",
            ),
            Expectation("no_final_signoff", r"you should sign|final legal advice|sign now", should_match=False),
        ),
    ),
    # Output structure
    EvalCase(
        "O-01",
        "output",
        "high",
        (
            "Review this term sheet excerpt and include the following in your output: "
            "risk dashboard, negotiation priorities, confidence score, and one-line legal boundary.\n"
            "Excerpt:\n"
            "- Series A preferred stock, $5M at $25M post-money\n"
            "- Full ratchet anti-dilution\n"
            "- 2x participating liquidation preference\n"
            "- No board observer seat for Paradigm\n"
            "- No amendment veto for Paradigm"
        ),
        (
            Expectation("risk_dashboard", r"risk|dashboard|RED_LINE|severity"),
            Expectation("negotiation_priorities", r"negotiat|priorit|must.fix|action"),
            Expectation("confidence", r"confidence"),
            Expectation(
                "disclaimer",
                r"not (a|your) lawyer|not (a )?licensed attorney|not legal advice",
            ),
        ),
    ),
)


CASE_TIMEOUT_OVERRIDES_S: dict[str, int] = {
    "R-01": 90,
    "R-02": 240,
    "N-01": 210,
    "N-02": 240,
    "N-03": 240,
    "T-04": 210,
    "K-04": 240,
}


def _load_dotenv(path: Path | None) -> dict[str, str]:
    try:
        from dotenv import dotenv_values
    except Exception as exc:  # pragma: no cover - import guard
        raise RuntimeError(
            "python-dotenv is required. Run with: uv run --with python-dotenv ..."
        ) from exc

    if path is None:
        return {}
    values = dotenv_values(str(path))
    env: dict[str, str] = {}
    for key, value in values.items():
        if value is not None:
            env[key] = str(value).strip()
    return env


def _find_env_file(explicit: str | None) -> Path | None:
    if explicit:
        path = Path(explicit).expanduser().resolve()
        if not path.exists():
            raise FileNotFoundError(f"Explicit --env-file path not found: {path}")
        return path

    candidates = [
        Path.cwd() / ".env",
        Path.cwd().parent / ".env",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def _tier_rank(tier: str) -> int:
    return {"critical": 0, "high": 1, "extended": 2}[tier]


def _filter_cases(
    tier: str,
    categories: set[str],
    case_ids: set[str],
    limit: int | None,
) -> list[EvalCase]:
    max_rank = {"critical": 0, "high": 1, "all": 2}[tier]
    selected = [c for c in CASES if _tier_rank(c.tier) <= max_rank]
    if categories:
        selected = [c for c in selected if c.category in categories]
    if case_ids:
        selected = [c for c in selected if c.case_id in case_ids]
    if limit is not None:
        selected = selected[:limit]
    return selected


def _regex_match(text: str, exp: Expectation) -> bool:
    normalized = (
        text.replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u2019", "'")
    )
    matched = re.search(exp.pattern, normalized, re.IGNORECASE | re.MULTILINE) is not None
    return matched if exp.should_match else (not matched)


def _setup_agent_client(env: dict[str, str]) -> Any:
    try:
        import api.agent as agent_mod
    except Exception as exc:  # pragma: no cover - import guard
        raise RuntimeError(
            "Failed importing api.agent. Run with: uv run --with psycopg2-binary ..."
        ) from exc

    def val(key: str) -> str:
        return env.get(key) or os.getenv(key, "")

    def local_env() -> list[str]:
        api_key = val("AI_V2_API_KEY") or val("API_SECRET_KEY")
        out = [
            f"AI_V2_API_URL={val('AI_V2_API_URL') or 'http://host.docker.internal:8000'}",
            f"AI_V2_API_KEY={api_key}",
        ]
        for key in (
            "ANTHROPIC_API_KEY",
            "OPENAI_API_KEY",
            "CODEX_API_KEY",
            "AMP_API_KEY",
            "AMPCODE_API_KEY",
            "GITHUB_TOKEN",
        ):
            value = val(key)
            if value:
                out.append(f"{key}={value}")
        return out

    agent_mod._container_env = local_env
    return agent_mod.AgentClient()


def _validate_runner_env(env: dict[str, str]) -> None:
    api_key = (
        env.get("AI_V2_API_KEY")
        or os.getenv("AI_V2_API_KEY")
        or env.get("API_SECRET_KEY")
        or os.getenv("API_SECRET_KEY")
    )
    missing: list[str] = []
    if not api_key:
        missing.append("AI_V2_API_KEY (or API_SECRET_KEY)")
    if missing:
        raise RuntimeError(f"Missing required env vars for eval runner: {', '.join(missing)}")


def _effective_case_timeout(case_id: str, default_timeout_s: int) -> int:
    return max(default_timeout_s, CASE_TIMEOUT_OVERRIDES_S.get(case_id, default_timeout_s))


def _run_docx_fidelity_eval() -> list[dict[str, Any]]:
    try:
        from tools.termsheet.client import TermsheetClient
        from tools.termsheet.models import (
            BoardRights,
            InstrumentType,
            TermIntent,
            TermSheet,
            TokenRights,
        )
    except Exception as exc:
        return [
            {
                "case_id": "D-00",
                "category": "docx",
                "passed": False,
                "error": f"DOCX eval import failed: {exc}",
            }
        ]

    template_path = Path("tools/termsheet/templates/paradigm_term_sheet.docx").resolve()
    if not template_path.exists():
        return [
            {
                "case_id": "D-00",
                "category": "docx",
                "passed": False,
                "error": f"Template not found: {template_path}",
            }
        ]

    client = TermsheetClient()

    def make_doc(ts: TermSheet) -> bytes:
        return client.generate_docx(ts, template_file=str(template_path))

    def inspect_docx(docx_path: Path | None, docx_bytes: bytes | None) -> dict[str, Any]:
        data = docx_bytes if docx_bytes is not None else docx_path.read_bytes()  # type: ignore[arg-type]

        with zipfile.ZipFile(io.BytesIO(data), "r") as zf:
            names = zf.namelist()
            headers = sorted(n for n in names if n.startswith("word/header") and n.endswith(".xml"))
            media = sorted(n for n in names if n.startswith("word/media/"))
            header_rels = sorted(
                n for n in names if n.startswith("word/_rels/header") and n.endswith(".rels")
            )
            xml_blob = ""
            for name in names:
                if name.startswith("word/") and name.endswith(".xml"):
                    xml_blob += zf.read(name).decode("utf-8", errors="ignore")
            unresolved_markers = [
                marker
                for marker in ("{{", "}}", "[__]", "[COMPANY]", "$[__]M")
                if marker in xml_blob
            ]
            media_hashes = {name: hashlib.sha256(zf.read(name)).hexdigest() for name in media}
            metadata = {
                "headers": headers,
                "header_rels": header_rels,
                "media": media,
                "media_hashes": media_hashes,
                "unresolved_markers": unresolved_markers,
            }
        doc = Document(io.BytesIO(data))
        text = "\n".join(
            [p.text for p in doc.paragraphs]
            + [
                p.text
                for table in doc.tables
                for row in table.rows
                for cell in row.cells
                for p in cell.paragraphs
            ]
        )
        metadata["text"] = text
        return metadata

    template_info = inspect_docx(template_path, None)

    baseline = client.create_term_sheet(
        company_name="Acme Robotics",
        instrument_type=InstrumentType.PRICED,
        investment_amount=6_300_000,
        post_money_valuation=21_000_000,
        option_pool_percent=10,
        board_rights=BoardRights.OBSERVER,
        legal_fee_cap=100_000,
        nvca_year=2025,
        token_rights=TokenRights(enabled=True, token_floor_percent=50),
        debt_threshold=1_000_000,
        founder_carveout_percent=2.0,
    )
    baseline_info = inspect_docx(None, make_doc(baseline))

    high_round = client.create_term_sheet(
        company_name="LargeCo",
        instrument_type=InstrumentType.PRICED,
        investment_amount=50_000_000,
        post_money_valuation=500_000_000,
        series="A",
        debt_threshold=5_000_000,
        token_rights=TokenRights(enabled=True, token_floor_percent=50),
    )
    high_round_info = inspect_docx(None, make_doc(high_round))

    seed_default = client.create_term_sheet(
        company_name="SeedCo",
        instrument_type=InstrumentType.PRICED,
        investment_amount=6_000_000,
        post_money_valuation=30_000_000,
        series="Seed",
        intent=TermIntent.FOUNDER_FRIENDLY,
        token_rights=TokenRights(enabled=True, token_floor_percent=50),
    )
    seed_info = inspect_docx(None, make_doc(seed_default))

    no_token = client.create_term_sheet(
        company_name="NoTokenCo",
        instrument_type=InstrumentType.PRICED,
        investment_amount=8_000_000,
        post_money_valuation=40_000_000,
        token_rights=TokenRights(enabled=False),
    )
    no_token_info = inspect_docx(None, make_doc(no_token))

    seat_only = client.create_term_sheet(
        company_name="SeatOnlyCo",
        instrument_type=InstrumentType.PRICED,
        investment_amount=10_000_000,
        post_money_valuation=100_000_000,
        board_rights=BoardRights.SEAT,
        token_rights=TokenRights(enabled=False),
    )
    seat_only_info = inspect_docx(None, make_doc(seat_only))

    custom_override = client.create_term_sheet(
        company_name="CustomCo",
        instrument_type=InstrumentType.PRICED,
        investment_amount=10_000_000,
        post_money_valuation=80_000_000,
        token_rights=TokenRights(enabled=True, token_floor_percent=55),
        protective_provision_v_text="any interested or related party transactions subject to customary exceptions including employee benefits and board-approved grants",
        vesting_text="Founder vesting to begin on date founders started working on the project.",
    )
    custom_info = inspect_docx(None, make_doc(custom_override))

    results: list[dict[str, Any]] = []
    results.append(
        {
            "case_id": "D-01",
            "category": "docx",
            "passed": template_info["headers"] == baseline_info["headers"],
            "checks": {
                "template_headers": template_info["headers"],
                "generated_headers": baseline_info["headers"],
            },
        }
    )
    results.append(
        {
            "case_id": "D-02",
            "category": "docx",
            "passed": template_info["media_hashes"] == baseline_info["media_hashes"],
            "checks": {
                "template_media_count": len(template_info["media_hashes"]),
                "generated_media_count": len(baseline_info["media_hashes"]),
            },
        }
    )
    results.append(
        {
            "case_id": "D-03",
            "category": "docx",
            "passed": not baseline_info["unresolved_markers"],
            "checks": {"unresolved_markers": baseline_info["unresolved_markers"]},
        }
    )
    results.append(
        {
            "case_id": "D-04",
            "category": "docx",
            "passed": "ACME ROBOTICS" in baseline_info["text"],
            "checks": {"company_caps_present": "ACME ROBOTICS" in baseline_info["text"]},
        }
    )
    results.append(
        {
            "case_id": "D-05",
            "category": "docx",
            "passed": "$5M" in high_round_info["text"],
            "checks": {"debt_threshold_5m_present": "$5M" in high_round_info["text"]},
        }
    )
    results.append(
        {
            "case_id": "D-06",
            "category": "docx",
            "passed": "2025 NVCA forms" in baseline_info["text"]
            and "up to $100,000" in baseline_info["text"],
            "checks": {
                "nvca_2025_present": "2025 NVCA forms" in baseline_info["text"],
                "fee_cap_100k_present": "up to $100,000" in baseline_info["text"],
            },
        }
    )
    results.append(
        {
            "case_id": "D-07",
            "category": "docx",
            "passed": "net proceeds greater than $50M" in seed_info["text"]
            and "together with other series of Preferred Stock" not in seed_info["text"],
            "checks": {
                "seed_ipo_50m": "net proceeds greater than $50M" in seed_info["text"],
                "seed_series_clause_trimmed": "together with other series of Preferred Stock"
                not in seed_info["text"],
            },
        }
    )
    results.append(
        {
            "case_id": "D-08",
            "category": "docx",
            "passed": bool(
                re.search(
                    r"up to\s+5(?:\.0+)?%\s+of the stock initially",
                    seed_info["text"],
                    re.IGNORECASE,
                )
            ),
            "checks": {
                "founder_carveout_5pct": bool(
                    re.search(
                        r"up to\s+5(?:\.0+)?%\s+of the stock initially",
                        seed_info["text"],
                        re.IGNORECASE,
                    )
                )
            },
        }
    )
    results.append(
        {
            "case_id": "D-09",
            "category": "docx",
            "passed": "Token Rights:" not in no_token_info["text"],
            "checks": {"token_rights_removed": "Token Rights:" not in no_token_info["text"]},
        }
    )
    results.append(
        {
            "case_id": "D-10",
            "category": "docx",
            "passed": "One director to be elected by the Series" in seat_only_info["text"]
            and "nonvoting observer capacity" not in seat_only_info["text"],
            "checks": {
                "seat_clause_present": "One director to be elected by the Series"
                in seat_only_info["text"],
                "observer_clause_absent": "nonvoting observer capacity"
                not in seat_only_info["text"],
            },
        }
    )
    results.append(
        {
            "case_id": "D-11",
            "category": "docx",
            "passed": all(
                snippet in baseline_info["text"]
                for snippet in (
                    "exclusive of granted or promised shares",
                    "Bylaws to provide for transfer restrictions on Common Stock",
                    "Customary closing conditions",
                )
            ),
            "checks": {
                "required_snippets_present": all(
                    snippet in baseline_info["text"]
                    for snippet in (
                        "exclusive of granted or promised shares",
                        "Bylaws to provide for transfer restrictions on Common Stock",
                        "Customary closing conditions",
                    )
                )
            },
        }
    )
    results.append(
        {
            "case_id": "D-12",
            "category": "docx",
            "passed": "subject to customary exceptions including employee benefits and board-approved grants."
            in custom_info["text"]
            and "Founder vesting to begin on date founders started working on the project."
            in custom_info["text"],
            "checks": {
                "custom_pp_v_present": "subject to customary exceptions including employee benefits and board-approved grants."
                in custom_info["text"],
                "custom_vesting_present": "Founder vesting to begin on date founders started working on the project."
                in custom_info["text"],
            },
        }
    )

    manifest_exists = False
    with tempfile.TemporaryDirectory(prefix="termsheet-delivery-") as tmp_dir:
        delivery_artifacts = client.generate_document_package(
            baseline,
            output_dir=tmp_dir,
            include_pdf=False,
            write_manifest=True,
            slack_channel="#deal-closing",
            slack_thread_ts="1234567890.000001",
        )
        manifest_path = Path(str(delivery_artifacts.get("delivery_manifest_json", "")))
        delivery_manifest: dict[str, Any] = {}
        manifest_exists = manifest_path.exists()
        if manifest_exists:
            delivery_manifest = json.loads(manifest_path.read_text())

    fidelity = delivery_manifest.get("fidelity", {})
    slack_delivery = delivery_manifest.get("slack_delivery", {})
    results.append(
        {
            "case_id": "D-13",
            "category": "docx",
            "passed": manifest_exists and bool(fidelity.get("passed")),
            "checks": {
                "manifest_exists": manifest_exists,
                "fidelity_passed": bool(fidelity.get("passed")),
            },
        }
    )
    results.append(
        {
            "case_id": "D-14",
            "category": "docx",
            "passed": bool(fidelity.get("banner_integrity"))
            and bool(fidelity.get("fonts", {}).get("unchanged"))
            and bool(fidelity.get("style_ids", {}).get("unchanged"))
            and bool(fidelity.get("protected_parts_present"))
            and bool(fidelity.get("protected_parts_unchanged"))
            and bool(fidelity.get("header_parts_unchanged"))
            and bool(fidelity.get("header_rel_parts_unchanged")),
            "checks": {
                "banner_integrity": bool(fidelity.get("banner_integrity")),
                "fonts_unchanged": bool(fidelity.get("fonts", {}).get("unchanged")),
                "style_ids_unchanged": bool(fidelity.get("style_ids", {}).get("unchanged")),
                "protected_parts_present": bool(fidelity.get("protected_parts_present")),
                "protected_parts_unchanged": bool(fidelity.get("protected_parts_unchanged")),
                "header_parts_unchanged": bool(fidelity.get("header_parts_unchanged")),
                "header_rel_parts_unchanged": bool(fidelity.get("header_rel_parts_unchanged")),
            },
        }
    )
    results.append(
        {
            "case_id": "D-15",
            "category": "docx",
            "passed": bool(slack_delivery.get("all_sendable"))
            and bool(delivery_manifest.get("delivery_ready")),
            "checks": {
                "all_sendable": bool(slack_delivery.get("all_sendable")),
                "delivery_ready": bool(delivery_manifest.get("delivery_ready")),
                "files_checked": len(slack_delivery.get("files", [])),
            },
        }
    )
    return results


def _run_mode_guard_eval() -> list[dict[str, Any]]:
    legal_mode_path = Path("apps/slackbot/src/lib/modes/legal.ts").resolve()
    bot_mode_path = Path("apps/slackbot/src/lib/bot.ts").resolve()
    if not legal_mode_path.exists():
        return [
            {
                "case_id": "OP-00",
                "category": "ops",
                "passed": False,
                "error": f"legal mode file not found: {legal_mode_path}",
            }
        ]
    if not bot_mode_path.exists():
        return [
            {
                "case_id": "OP-00B",
                "category": "ops",
                "passed": False,
                "error": f"bot mode file not found: {bot_mode_path}",
            }
        ]

    text = legal_mode_path.read_text()
    bot_text = bot_mode_path.read_text()
    continue_true = "continueSession: true" in text
    continue_false = "continueSession: false" in text
    request_id_scoped = "requestId: `${requestId}:${phaseLabel}`" in text
    term_sheet_hinting = (
        "call termsheet create_term_sheet" in text
        and "call termsheet explain_clause_plan" in text
        and "call termsheet generate_document_package" in text
    )
    retrieval_hinting = "call search" in text
    compliance_tool_hinting = (
        "call legal-playbook check_compliance" in text
        and "call legal-playbook score_quality" in text
    )
    legal_kickoff_unified = (
        "parsed.cleanedText || buildLegalKickoffInstruction()" in bot_text
        and "if (!parsed.cleanedText && isLegalHarness(harness))" not in bot_text
    )
    legal_identity_phrase = (
        "I am not a lawyer; I am a legal agent created by Paradigm." in text
    )
    legal_kickoff_contract = (
        "provide exactly 4 short example prompts" in text
        and "ask one focused follow-up question to begin work" in text
        and "keep output <= 220 words and avoid marketing tone" in text
    )
    lawyer_style_contract = (
        "Style contract for this phase:" in text
        and "Write like a senior transactional lawyer" in text
    )

    return [
        {
            "case_id": "OP-01",
            "category": "ops",
            "passed": continue_true and not continue_false,
            "checks": {
                "continue_true_present": continue_true,
                "continue_false_absent": not continue_false,
            },
        },
        {
            "case_id": "OP-02",
            "category": "ops",
            "passed": request_id_scoped,
            "checks": {"phase_scoped_request_ids": request_id_scoped},
        },
        {
            "case_id": "OP-03",
            "category": "ops",
            "passed": term_sheet_hinting,
            "checks": {"tool_grounded_term_sheet_prompts": term_sheet_hinting},
        },
        {
            "case_id": "OP-04",
            "category": "ops",
            "passed": retrieval_hinting,
            "checks": {"shared_search_retrieval_prompts": retrieval_hinting},
        },
        {
            "case_id": "OP-05",
            "category": "ops",
            "passed": compliance_tool_hinting,
            "checks": {"tool_backed_compliance_prompts": compliance_tool_hinting},
        },
        {
            "case_id": "OP-06",
            "category": "ops",
            "passed": legal_kickoff_unified and legal_identity_phrase and legal_kickoff_contract,
            "checks": {
                "legal_kickoff_unified_path": legal_kickoff_unified,
                "legal_identity_phrase_present": legal_identity_phrase,
                "legal_kickoff_contract_present": legal_kickoff_contract,
            },
        },
        {
            "case_id": "OP-07",
            "category": "ops",
            "passed": lawyer_style_contract,
            "checks": {"lawyer_style_contract_present": lawyer_style_contract},
        },
    ]


def _render_markdown(report: dict[str, Any]) -> str:
    lines: list[str] = []
    lines.append("# Legal Eval Report")
    lines.append("")
    lines.append(f"- Timestamp: `{report['timestamp']}`")
    lines.append(f"- Model: `{report['model']}`")
    lines.append(f"- Cases run: `{report['total_cases']}`")
    lines.append(f"- Passed: `{report['passed_cases']}`")
    lines.append(f"- Failed: `{report['failed_cases']}`")
    lines.append(f"- Pass rate: `{report['pass_rate']:.1f}%`")
    artifacts_root = report.get("artifacts_root")
    if artifacts_root:
        lines.append(f"- Case artifacts: `{artifacts_root}`")
    lines.append("")
    lines.append("## Category Summary")
    lines.append("")
    lines.append("| Category | Passed | Total | Pass % |")
    lines.append("|---|---:|---:|---:|")
    for category, stats in sorted(report["by_category"].items()):
        lines.append(
            f"| {category} | {stats['passed']} | {stats['total']} | {stats['pass_rate']:.1f}% |"
        )
    lines.append("")
    lines.append("## Case Artifacts")
    lines.append("")
    lines.append("| Case | Category | Pass | Prompt | Output |")
    lines.append("|---|---|---:|---|---|")
    for item in report["results"]:
        prompt_path = item.get("prompt_path", "")
        output_path = item.get("output_path", "")
        lines.append(
            "| "
            f"{item.get('case_id', 'n/a')} | "
            f"{item.get('category', 'n/a')} | "
            f"{'✅' if item.get('passed', False) else '❌'} | "
            f"`{prompt_path}` | "
            f"`{output_path}` |"
        )
    lines.append("")
    lines.append("## Failed Cases")
    lines.append("")
    failed = [r for r in report["results"] if not r.get("passed", False)]
    if not failed:
        lines.append("- None")
    else:
        for item in failed:
            lines.append(f"### {item['case_id']} ({item.get('category', 'n/a')})")
            if item.get("error"):
                lines.append(f"- Error: `{item['error']}`")
            if item.get("expectations"):
                for exp in item["expectations"]:
                    if not exp["passed"]:
                        lines.append(f"- Failed check `{exp['name']}`")
            preview = str(item.get("preview", "")).strip()
            if preview:
                lines.append(f"- Preview: `{preview[:360]}`")
            lines.append("")
    return "\n".join(lines) + "\n"


def _write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


def _sha256_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


class EvalCaseTimeoutError(Exception):
    pass


def _execute_case_worker(
    result_path: str,
    env: dict[str, str],
    case_thread_key: str,
    prompt: str,
    model: str,
) -> None:
    client: Any | None = None
    payload: dict[str, Any]
    try:
        client = _setup_agent_client(env)
        response = client.execute(
            case_thread_key,
            prompt,
            harness="legal",
            source="api",
            model=model,
            continue_session=False,
        )
        payload = {"ok": True, "response": dict(response)}
    except Exception as exc:
        payload = {"ok": False, "error": str(exc), "traceback": traceback.format_exc()}
    finally:
        path = Path(result_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(json.dumps(payload, default=str))


def _force_remove_case_container(case_thread_key: str) -> None:
    """Best-effort cleanup for case-scoped sandbox containers."""
    try:
        import docker
    except Exception:
        return

    with contextlib.suppress(Exception):
        client = docker.from_env()
        containers = client.containers.list(
            all=True,
            filters={"label": [f"ai2.thread={case_thread_key}", "agent2=true"]},
        )
        for container in containers:
            with contextlib.suppress(Exception):
                container.stop(timeout=3)
            with contextlib.suppress(Exception):
                container.remove(force=True)


def main() -> int:
    parser = argparse.ArgumentParser(description="Run Ben-aligned legal eval suite.")
    parser.add_argument("--model", default="sonnet", help="Model alias passed to harness.")
    parser.add_argument(
        "--tier",
        default="high",
        choices=("critical", "high", "all"),
        help="Case tier selection.",
    )
    parser.add_argument("--category", action="append", default=[], help="Case category filter.")
    parser.add_argument("--case-id", action="append", default=[], help="Specific case IDs.")
    parser.add_argument("--limit", type=int, default=None, help="Limit number of prompt cases.")
    parser.add_argument("--env-file", default=None, help="Optional .env file path.")
    parser.add_argument("--output-dir", default="evals/legal", help="Report output directory.")
    parser.add_argument("--skip-docx", action="store_true", help="Skip DOCX fidelity evals.")
    parser.add_argument(
        "--capture-case-files",
        action=argparse.BooleanOptionalAction,
        default=True,
        help="Write exact prompt/output files per case (default: on).",
    )
    parser.add_argument(
        "--include-full-output-json",
        action=argparse.BooleanOptionalAction,
        default=False,
        help="Embed full prompt/output text in the report JSON (default: off).",
    )
    parser.add_argument(
        "--case-timeout-s",
        type=int,
        default=180,
        help="Per-case hard timeout in seconds (0 disables).",
    )
    args = parser.parse_args()

    root = Path(__file__).resolve().parents[1]
    os.chdir(root)
    sys.path.insert(0, str(root))
    sys.path.insert(0, str(root / "src"))

    try:
        env_file = _find_env_file(args.env_file)
    except FileNotFoundError as exc:
        print(str(exc))
        return 1
    env = dict(os.environ)
    env.update(_load_dotenv(env_file))
    _validate_runner_env(env)

    selected = _filter_cases(
        tier=args.tier,
        categories={c.strip() for c in args.category if c.strip()},
        case_ids={c.strip() for c in args.case_id if c.strip()},
        limit=args.limit,
    )
    if not selected and args.skip_docx:
        print("No cases selected.")
        return 1

    print(f"Selected prompt cases: {len(selected)}")
    if env_file:
        print(f"Using env file: {env_file}")

    client = _setup_agent_client(env) if selected else None
    timestamp = datetime.now(UTC).strftime("%Y%m%dT%H%M%SZ")
    results: list[dict[str, Any]] = []
    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    artifacts_root = out_dir / f"legal-eval-{timestamp}-cases"

    shared_run_key = f"eval:runner:{timestamp}"

    for idx, case in enumerate(selected):
        print(f"\n=== [{idx + 1}/{len(selected)}] {case.case_id} ({case.category}) ===")
        started = time.time()
        case_thread_key = f"{shared_run_key}:{case.case_id}"
        case_dir = artifacts_root / case.case_id
        prompt_path = case_dir / "prompt.txt"
        output_path = case_dir / "output.txt"
        response_path = case_dir / "response.json"
        error_path = case_dir / "error.txt"
        if args.capture_case_files:
            _write_text(prompt_path, case.prompt + "\n")
        try:
            response: dict[str, Any]
            timeout_s = _effective_case_timeout(case.case_id, args.case_timeout_s)
            if args.case_timeout_s > 0:
                case_dir.mkdir(parents=True, exist_ok=True)
                worker_result_path = case_dir / "worker_result.json"
                if worker_result_path.exists():
                    worker_result_path.unlink()
                ctx = multiprocessing.get_context("spawn")
                process = ctx.Process(
                    target=_execute_case_worker,
                    args=(
                        str(worker_result_path),
                        env,
                        case_thread_key,
                        case.prompt,
                        args.model,
                    ),
                )
                process.start()
                deadline = time.monotonic() + timeout_s
                while process.is_alive():
                    remaining = deadline - time.monotonic()
                    if remaining <= 0:
                        process.terminate()
                        process.join()
                        _force_remove_case_container(case_thread_key)
                        raise EvalCaseTimeoutError(f"case execution timed out (timeout_s={timeout_s})")
                    process.join(min(1.0, remaining))
                if not worker_result_path.exists():
                    raise RuntimeError("case worker exited without writing a result")
                worker_result = json.loads(worker_result_path.read_text())
                if not worker_result.get("ok"):
                    err = worker_result.get("error", "unknown case worker failure")
                    tb = worker_result.get("traceback")
                    raise RuntimeError(f"{err}\n{tb}" if tb else err)
                response = dict(worker_result["response"])
            else:
                if client is None:
                    raise RuntimeError("agent client is not initialized")
                response = client.execute(
                    case_thread_key,
                    case.prompt,
                    harness="legal",
                    source="api",
                    model=args.model,
                    continue_session=False,
                )
            if response.get("error"):
                raise RuntimeError(f"agent returned error: {response['error']}")
            text = str(response.get("result", ""))
            expectation_results = []
            for exp in case.expectations:
                passed = _regex_match(text, exp)
                expectation_results.append(
                    {"name": exp.name, "passed": passed, "pattern": exp.pattern}
                )
            expectation_results.append(
                {
                    "name": "non_empty_output",
                    "passed": bool(text.strip()),
                    "pattern": r"\S+",
                }
            )
            if case.max_chars is not None:
                expectation_results.append(
                    {
                        "name": "max_chars",
                        "passed": len(text) <= case.max_chars,
                        "pattern": f"len(output)<= {case.max_chars}",
                    }
                )
            passed = all(item["passed"] for item in expectation_results)
            if args.capture_case_files:
                _write_text(output_path, text + ("\n" if not text.endswith("\n") else ""))
                _write_text(response_path, json.dumps(response, indent=2, default=str))
            results.append(
                {
                    "case_id": case.case_id,
                    "category": case.category,
                    "tier": case.tier,
                    "passed": passed,
                    "duration_s": round(time.time() - started, 3),
                    "harness": response.get("harness"),
                    "engine": response.get("engine"),
                    "persona": response.get("persona"),
                    "chars": len(text),
                    "prompt_sha256": _sha256_text(case.prompt),
                    "output_sha256": _sha256_text(text),
                    "expectations": expectation_results,
                    "preview": " ".join(text.split())[:700],
                    "prompt_path": str(prompt_path) if args.capture_case_files else None,
                    "output_path": str(output_path) if args.capture_case_files else None,
                    "response_path": str(response_path) if args.capture_case_files else None,
                }
            )
            if args.include_full_output_json:
                results[-1]["prompt_text"] = case.prompt
                results[-1]["output_text"] = text
            print(f"passed={passed} chars={len(text)}")
        except Exception as exc:
            if args.capture_case_files:
                _write_text(error_path, f"{exc}\n")
            results.append(
                {
                    "case_id": case.case_id,
                    "category": case.category,
                    "tier": case.tier,
                    "passed": False,
                    "duration_s": round(time.time() - started, 3),
                    "error": str(exc),
                    "prompt_path": str(prompt_path) if args.capture_case_files else None,
                    "output_path": str(output_path) if args.capture_case_files else None,
                    "response_path": str(response_path) if args.capture_case_files else None,
                    "error_path": str(error_path) if args.capture_case_files else None,
                }
            )
            if args.include_full_output_json:
                results[-1]["prompt_text"] = case.prompt
            print(f"failed with exception: {exc}")
        finally:
            if client is not None:
                with contextlib.suppress(Exception):
                    client.stop(case_thread_key)
            _force_remove_case_container(case_thread_key)

    if client is not None:
        with contextlib.suppress(Exception):
            client.stop(shared_run_key)

    if not args.skip_docx:
        print("\n=== DOCX fidelity checks ===")
        results.extend(_run_docx_fidelity_eval())
    print("\n=== Legal mode guard checks ===")
    results.extend(_run_mode_guard_eval())

    total_cases = len(results)
    passed_cases = sum(1 for r in results if r.get("passed", False))
    failed_cases = total_cases - passed_cases
    pass_rate = (passed_cases / total_cases * 100.0) if total_cases else 0.0

    by_category: dict[str, dict[str, Any]] = {}
    for item in results:
        category = item.get("category", "unknown")
        stats = by_category.setdefault(category, {"passed": 0, "total": 0, "pass_rate": 0.0})
        stats["total"] += 1
        if item.get("passed", False):
            stats["passed"] += 1
    for stats in by_category.values():
        stats["pass_rate"] = (stats["passed"] / stats["total"] * 100.0) if stats["total"] else 0.0

    report = {
        "timestamp": timestamp,
        "model": args.model,
        "env_file": str(env_file) if env_file else None,
        "artifacts_root": str(artifacts_root) if args.capture_case_files else None,
        "total_cases": total_cases,
        "passed_cases": passed_cases,
        "failed_cases": failed_cases,
        "pass_rate": pass_rate,
        "by_category": by_category,
        "results": results,
    }

    json_path = out_dir / f"legal-eval-{timestamp}.json"
    md_path = out_dir / f"legal-eval-{timestamp}.md"
    json_path.write_text(json.dumps(report, indent=2))
    md_path.write_text(_render_markdown(report))

    print("\n=== Summary ===")
    print(f"pass_rate={pass_rate:.1f}% ({passed_cases}/{total_cases})")
    print(f"json={json_path}")
    print(f"md={md_path}")
    return 0 if failed_cases == 0 else 2


if __name__ == "__main__":
    raise SystemExit(main())
